#!/usr/bin/env python
# -*- coding: UTF-8 -*-
# Copyright (c) Huawei Technologies Co., Ltd. 2026. All rights reserved.

"""Unit tests for FileParser utility class (helpers + parsers)."""

import io
from unittest.mock import patch

import pytest

from agent_runtime.utils.file_parser import FileParser


# ---------------------------------------------------------------------------
# Helpers to create test files in memory
# ---------------------------------------------------------------------------


def _make_xlsx_bytes():
    """Create an xlsx file in memory with 2 sheets."""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sheet1"
    ws1.append(["name", "age"])
    ws1.append(["Alice", 30])
    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["city", "pop"])
    ws2.append(["NYC", "8M"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_docx_bytes():
    """Create a docx file in memory with interleaved paragraphs and table."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Before Table")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "C"
    table.rows[1].cells[1].text = "D"
    doc.add_paragraph("After Table")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# build_legal_name
# ===========================================================================


class TestBuildLegalName:
    """Tests for FileParser.build_legal_name()."""

    @staticmethod
    def test_valid_name_returned_as_is():
        assert FileParser.build_legal_name("report.txt") == "report.txt"

    @staticmethod
    def test_empty_string_returns_uuid():
        result = FileParser.build_legal_name("")
        assert len(result) == 36
        assert result.count("-") == 4

    @staticmethod
    def test_none_returns_uuid():
        result = FileParser.build_legal_name(None)
        assert len(result) == 36

    @pytest.mark.parametrize("name", ["a/b", "c:d", "e*f", 'g"h', "i<j", "j>k", "k|l", "m\\n", "n?o"])
    @staticmethod
    def test_illegal_chars_return_uuid(name):
        result = FileParser.build_legal_name(name)
        assert len(result) == 36, f"Expected UUID for {name!r}, got {result!r}"

    @staticmethod
    def test_space_is_legal():
        assert FileParser.build_legal_name("my report.txt") == "my report.txt"


# ===========================================================================
# get_suffix
# ===========================================================================


class TestGetSuffix:
    """Tests for FileParser.get_suffix()."""

    @staticmethod
    def test_simple_txt():
        assert FileParser.get_suffix("https://example.com/file.txt") == "txt"

    @staticmethod
    def test_uppercase_extension_lowered():
        assert FileParser.get_suffix("https://example.com/file.TXT") == "txt"

    @staticmethod
    def test_no_extension_returns_empty():
        assert FileParser.get_suffix("https://example.com/file") == ""

    @staticmethod
    def test_query_params_ignored():
        assert FileParser.get_suffix("https://example.com/file.csv?token=abc&expires=123") == "csv"

    @staticmethod
    def test_fragment_ignored():
        assert FileParser.get_suffix("https://example.com/file.docx#section") == "docx"

    @staticmethod
    def test_domain_dots_not_mistaken_for_extension():
        """https://example.com/noext should return '', not 'com/noext'."""
        assert FileParser.get_suffix("https://example.com/noext") == ""

    @staticmethod
    def test_multiple_dots_in_filename():
        assert FileParser.get_suffix("https://example.com/archive.tar.gz") == "gz"

    @staticmethod
    def test_path_with_subdirectories():
        assert FileParser.get_suffix("https://obs.example.com/bucket/path/to/data.xlsx") == "xlsx"

    @staticmethod
    def test_root_path_returns_empty():
        assert FileParser.get_suffix("https://example.com/") == ""

    @staticmethod
    def test_no_path_returns_empty():
        assert FileParser.get_suffix("https://example.com") == ""


# ===========================================================================
# truncate
# ===========================================================================


class TestTruncate:
    """Tests for FileParser.truncate()."""

    @staticmethod
    def test_short_text_unchanged():
        assert FileParser.truncate("hello", 100) == "hello"

    @staticmethod
    def test_exact_length_unchanged():
        assert FileParser.truncate("hello", 5) == "hello"

    @staticmethod
    def test_long_text_truncated():
        text = "a" * 100
        assert FileParser.truncate(text, 10) == "a" * 10

    @staticmethod
    def test_zero_max_returns_empty():
        assert FileParser.truncate("hello", 0) == ""


# ===========================================================================
# parse (dispatch)
# ===========================================================================


class TestParse:
    """Tests for FileParser.parse() dispatch."""

    @staticmethod
    def test_txt():
        assert FileParser.parse(b"hello\nworld", "txt") == "hello\nworld"

    @staticmethod
    def test_csv():
        result = FileParser.parse(b"a,b,c\n1,2,3", "csv")
        assert "a,b,c" in result
        assert "1,2,3" in result

    @staticmethod
    def test_csv_quoted_fields():
        data = b'name,note\n"Smith, John","hello, world"'
        result = FileParser.parse(data, "csv")
        assert "Smith, John" in result

    @staticmethod
    def test_xlsx_multi_sheet():
        data = _make_xlsx_bytes()
        result = FileParser.parse(data, "xlsx")
        assert "name\tage" in result
        assert "Alice\t30" in result
        assert "city\tpop" in result
        assert "NYC\t8M" in result

    @staticmethod
    def test_xlsx_empty_cells():
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["a", None, "c"])
        buf = io.BytesIO()
        wb.save(buf)
        result = FileParser.parse(buf.getvalue(), "xlsx")
        assert "a\t\tc" in result

    @staticmethod
    def test_docx_paragraphs_and_table():
        data = _make_docx_bytes()
        result = FileParser.parse(data, "docx")
        lines = result.split("\n")
        # 验证段落与表格按文档原始顺序输出：Before Table → 表格 → After Table
        assert lines[0] == "Before Table"
        assert lines[1] == "A\tB"
        assert lines[2] == "C\tD"
        assert lines[3] == "After Table"

    @staticmethod
    def test_docx_empty_document():
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = FileParser.parse(buf.getvalue(), "docx")
        assert result == ""

    @staticmethod
    def test_doc_parses_via_antiword():
        """Test that doc files are parsed via antiword subprocess with temp file."""
        from unittest.mock import MagicMock

        mock_result = MagicMock()
        mock_result.returncode = 0
        mock_result.stdout = b"Test document content\nWith multiple lines"
        mock_result.stderr = b""

        with patch("subprocess.run", return_value=mock_result) as mock_run, \
             patch("os.unlink") as mock_unlink:
            result = FileParser.parse(b"fake doc data", "doc")

            # Verify antiword was called with a temp file path (not stdin)
            mock_run.assert_called_once()
            call_args = mock_run.call_args
            cmd = call_args[0][0]
            assert cmd[0].endswith("antiword") or "antiword" in cmd[0].lower()
            assert cmd[1:3] == ["-m", "UTF-8.txt"]
            assert cmd[3].endswith(".doc")  # temp file path
            assert "input" not in call_args[1]  # no stdin input
            assert call_args[1]["capture_output"] is True

            # Verify temp file cleanup
            mock_unlink.assert_called_once()
            assert result == "Test document content\nWith multiple lines"

    @staticmethod
    def test_doc_antiword_failure_raises():
        """Test that antiword failure raises ValueError."""
        from unittest.mock import MagicMock

        mock_result = MagicMock()
        mock_result.returncode = 1
        mock_result.stdout = b""
        mock_result.stderr = b"Error: not a valid doc file"

        with patch("subprocess.run", return_value=mock_result):
            with pytest.raises(ValueError) as exc_info:
                FileParser.parse(b"invalid data", "doc")
            assert "antiword 解析失败" in str(exc_info.value)
            assert "not a valid doc file" in str(exc_info.value)

    @staticmethod
    def test_unsupported_xls_returns_message():
        result = FileParser.parse(b"dummy", "xls")
        assert "不支持" in result
        assert ".xlsx" in result

    @staticmethod
    def test_unknown_format_returns_empty():
        assert FileParser.parse(b"dummy", "unknown") == ""

    @staticmethod
    def test_empty_suffix_returns_empty():
        assert FileParser.parse(b"dummy", "") == ""

    @staticmethod
    def test_invalid_xlsx_raises():
        with pytest.raises(Exception):
            FileParser.parse(b"not a real xlsx", "xlsx")
