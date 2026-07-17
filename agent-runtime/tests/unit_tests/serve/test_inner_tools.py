#!/usr/bin/env python
# -*- coding: UTF-8 -*-
# Copyright (c) Huawei Technologies Co., Ltd. 2026. All rights reserved.

"""Unit tests for inner tools API endpoints (file resolve + document create).

Pure function tests for FileParser are in tests/unit_tests/common/test_file_parser.py.
"""

import io
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi import FastAPI
from httpx import ASGITransport, AsyncClient

from agent_runtime.serve.apis.inner_tools import inner_tools_router

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

_RESOLVE_MODULE = "agent_runtime.serve.apis.inner_tools"


def _make_app():
    """Build a FastAPI app with only the inner_tools router."""
    app = FastAPI()
    app.include_router(inner_tools_router)
    return app


def _make_xlsx_bytes():
    """Create an xlsx file in memory with 2 sheets."""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sheet1"
    ws1.append(["name", "age"])
    ws1.append(["Alice", 30])
    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["city", "pop"])
    ws2.append(["NYC", "8M"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_docx_bytes():
    """Create a docx file in memory with paragraphs + table."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("Second paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "C"
    table.rows[1].cells[1].text = "D"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_mock_settings(
    max_resolve_size=5000,
    obs_staging_bucket="staging-bucket",
    obs_expire_time=7,
):
    """Build a mock settings object with object_storage config."""
    mock_settings = MagicMock()
    mock_settings.object_storage.max_resolve_size = max_resolve_size
    mock_settings.object_storage.staging_bucket = obs_staging_bucket
    mock_settings.object_storage.expire_time = obs_expire_time
    return mock_settings


# ===========================================================================
# resolve_file endpoint tests
# ===========================================================================


class TestResolveFile:
    """Tests for POST /v1/inner-tools/file/resolve."""

    @pytest.mark.asyncio
    async def test_ssrf_rejected_returns_400(self):
        """When illegal_url returns True, request should be rejected with 400."""
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=True),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "http://internal-ip/file.txt"},
                )
        assert resp.status_code == 400
        assert "非法文件URL" in resp.json()["detail"]

    @pytest.mark.asyncio
    async def test_download_failure_returns_400(self):
        """Download failure should return 400 with error detail."""
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(
                f"{_RESOLVE_MODULE}._download_file",
                side_effect=Exception("network error"),
            ),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.txt"},
                )
        assert resp.status_code == 400
        assert "下载文件失败" in resp.json()["detail"]

    @pytest.mark.asyncio
    async def test_unsupported_xls_format_returns_message(self):
        """.xls format should return a guidance message, not an error."""
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(f"{_RESOLVE_MODULE}._download_file", return_value=b"dummy"),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.xls"},
                )
        assert resp.status_code == 200
        assert "不支持" in resp.json()["content"]
        assert ".xlsx" in resp.json()["content"]

    @pytest.mark.asyncio
    async def test_unknown_format_returns_empty_content(self):
        """Unknown format (no parser) should return empty content, not error."""
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(f"{_RESOLVE_MODULE}._download_file", return_value=b"dummy"),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.unknown"},
                )
        assert resp.status_code == 200
        assert resp.json()["content"] == ""

    @pytest.mark.asyncio
    async def test_txt_success(self):
        """TXT file should be parsed and content returned."""
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(
                f"{_RESOLVE_MODULE}._download_file",
                return_value="hello world".encode("utf-8"),
            ),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.txt"},
                )
        assert resp.status_code == 200
        assert resp.json()["content"] == "hello world"

    @pytest.mark.asyncio
    async def test_csv_success(self):
        """CSV file should be parsed and content returned."""
        csv_bytes = b"name,age\nAlice,30"
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(f"{_RESOLVE_MODULE}._download_file", return_value=csv_bytes),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.csv"},
                )
        assert resp.status_code == 200
        content = resp.json()["content"]
        assert "name,age" in content
        assert "Alice,30" in content

    @pytest.mark.asyncio
    async def test_xlsx_success(self):
        """XLSX file should be parsed and content returned."""
        xlsx_bytes = _make_xlsx_bytes()
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(f"{_RESOLVE_MODULE}._download_file", return_value=xlsx_bytes),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.xlsx"},
                )
        assert resp.status_code == 200
        content = resp.json()["content"]
        assert "name\tage" in content
        assert "Alice\t30" in content

    @pytest.mark.asyncio
    async def test_docx_success(self):
        """DOCX file should be parsed and content returned."""
        docx_bytes = _make_docx_bytes()
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(f"{_RESOLVE_MODULE}._download_file", return_value=docx_bytes),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.docx"},
                )
        assert resp.status_code == 200
        content = resp.json()["content"]
        assert "Hello World" in content
        assert "A\tB" in content

    @pytest.mark.asyncio
    async def test_content_truncated_to_max_size(self):
        """Content exceeding max_resolve_size should be truncated."""
        long_text = "x" * 200
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(
                f"{_RESOLVE_MODULE}._download_file",
                return_value=long_text.encode("utf-8"),
            ),
            patch(
                f"{_RESOLVE_MODULE}.settings",
                _make_mock_settings(max_resolve_size=50),
            ),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.txt"},
                )
        assert resp.status_code == 200
        assert len(resp.json()["content"]) == 50

    @pytest.mark.asyncio
    async def test_parser_failure_returns_500(self):
        """Parser exception should return 500."""
        with (
            patch(f"{_RESOLVE_MODULE}.illegal_url", return_value=False),
            patch(
                f"{_RESOLVE_MODULE}._download_file",
                return_value=b"not a real xlsx",
            ),
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/file/resolve",
                    params={"file_url": "https://example.com/file.xlsx"},
                )
        assert resp.status_code == 500
        assert "解析文件失败" in resp.json()["detail"]


# ===========================================================================
# create_document endpoint tests
# ===========================================================================


class TestCreateDocument:
    """Tests for POST /v1/inner-tools/document/create."""

    @pytest.mark.asyncio
    async def test_success_returns_url(self):
        """Successful document creation should return a presigned URL."""
        mock_storage = AsyncMock()
        mock_storage.put_object_bytes = AsyncMock()
        # get_presigned_url is async — aioboto3's generate_presigned_url returns a coroutine
        mock_storage.get_presigned_url = AsyncMock(
            return_value="https://obs.example.com/signed-url"
        )

        with (
            patch(f"{_RESOLVE_MODULE}.S3StorageProvider") as mock_s3_cls,
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            mock_s3_cls.instance.return_value = mock_storage
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/document/create",
                    json={"input": "line1\nline2", "document_name": "report.txt"},
                )
        assert resp.status_code == 200
        assert resp.json()["url"] == "https://obs.example.com/signed-url"
        mock_storage.put_object_bytes.assert_awaited_once()
        mock_storage.get_presigned_url.assert_awaited_once()

    @pytest.mark.asyncio
    async def test_upload_failure_returns_500(self):
        """Upload failure should return 500."""
        mock_storage = AsyncMock()
        mock_storage.put_object_bytes = AsyncMock(
            side_effect=Exception("OBS quota exceeded")
        )

        with (
            patch(f"{_RESOLVE_MODULE}.S3StorageProvider") as mock_s3_cls,
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            mock_s3_cls.instance.return_value = mock_storage
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/document/create",
                    json={"input": "content", "document_name": "report.txt"},
                )
        assert resp.status_code == 500
        assert "上传文档失败" in resp.json()["detail"]

    @pytest.mark.asyncio
    async def test_presigned_url_failure_returns_500(self):
        """Presigned URL generation failure should return 500."""
        mock_storage = AsyncMock()
        mock_storage.put_object_bytes = AsyncMock()
        mock_storage.get_presigned_url = AsyncMock(
            side_effect=Exception("signing error")
        )

        with (
            patch(f"{_RESOLVE_MODULE}.S3StorageProvider") as mock_s3_cls,
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            mock_s3_cls.instance.return_value = mock_storage
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/document/create",
                    json={"input": "content", "document_name": "report.txt"},
                )
        assert resp.status_code == 500
        assert "生成下载链接失败" in resp.json()["detail"]

    @pytest.mark.asyncio
    async def test_custom_expires_used(self):
        """Custom expires value should be converted from days to seconds."""
        mock_storage = AsyncMock()
        mock_storage.put_object_bytes = AsyncMock()
        mock_storage.get_presigned_url = AsyncMock(
            return_value="https://obs.example.com/signed"
        )

        with (
            patch(f"{_RESOLVE_MODULE}.S3StorageProvider") as mock_s3_cls,
            patch(
                f"{_RESOLVE_MODULE}.settings",
                _make_mock_settings(obs_expire_time=7),
            ),
        ):
            mock_s3_cls.instance.return_value = mock_storage
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/document/create",
                    json={
                        "input": "content",
                        "document_name": "report.txt",
                        "expires": 3,
                    },
                )
        assert resp.status_code == 200
        # expires=3 days → 3 * 86400 = 259200 seconds
        call_args = mock_storage.get_presigned_url.call_args
        expires_seconds = call_args.args[1]
        assert expires_seconds == 3 * 86400

    @pytest.mark.asyncio
    async def test_default_expires_used_when_not_provided(self):
        """When expires is not provided, default from settings should be used."""
        mock_storage = AsyncMock()
        mock_storage.put_object_bytes = AsyncMock()
        mock_storage.get_presigned_url = AsyncMock(
            return_value="https://obs.example.com/signed"
        )

        with (
            patch(f"{_RESOLVE_MODULE}.S3StorageProvider") as mock_s3_cls,
            patch(
                f"{_RESOLVE_MODULE}.settings",
                _make_mock_settings(obs_expire_time=14),
            ),
        ):
            mock_s3_cls.instance.return_value = mock_storage
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/document/create",
                    json={"input": "content", "document_name": "report.txt"},
                )
        assert resp.status_code == 200
        # default 14 days → 14 * 86400 = 1209600 seconds
        call_args = mock_storage.get_presigned_url.call_args
        expires_seconds = call_args.args[1]
        assert expires_seconds == 14 * 86400

    @pytest.mark.asyncio
    async def test_invalid_filename_uses_uuid_in_object_key(self):
        """Filename with illegal chars should be replaced with UUID."""
        mock_storage = AsyncMock()
        mock_storage.put_object_bytes = AsyncMock()
        mock_storage.get_presigned_url = AsyncMock(
            return_value="https://obs.example.com/signed"
        )

        with (
            patch(f"{_RESOLVE_MODULE}.S3StorageProvider") as mock_s3_cls,
            patch(f"{_RESOLVE_MODULE}.settings", _make_mock_settings()),
        ):
            mock_s3_cls.instance.return_value = mock_storage
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/document/create",
                    json={"input": "content", "document_name": "bad/name"},
                )
        assert resp.status_code == 200
        # Verify object_key uses UUID (format: file/{uuid}/{uuid}.docx)
        call_args = mock_storage.put_object_bytes.call_args
        object_key = call_args.args[0]
        assert object_key.startswith("file/")
        assert object_key.endswith(".docx")
        # The filename part should be a UUID + .docx (36 + 5 = 41 chars)
        filename = object_key.rsplit("/", 1)[-1]
        assert len(filename) == 41

    @pytest.mark.asyncio
    async def test_staging_bucket_passed_to_storage(self):
        """obs_staging_bucket from settings should be passed to storage methods."""
        mock_storage = AsyncMock()
        mock_storage.put_object_bytes = AsyncMock()
        mock_storage.get_presigned_url = AsyncMock(
            return_value="https://obs.example.com/signed"
        )

        with (
            patch(f"{_RESOLVE_MODULE}.S3StorageProvider") as mock_s3_cls,
            patch(
                f"{_RESOLVE_MODULE}.settings",
                _make_mock_settings(obs_staging_bucket="my-staging-bucket"),
            ),
        ):
            mock_s3_cls.instance.return_value = mock_storage
            app = _make_app()
            async with AsyncClient(
                transport=ASGITransport(app=app), base_url="http://test"
            ) as client:
                resp = await client.post(
                    "/v1/inner-tools/document/create",
                    json={"input": "content", "document_name": "report.txt"},
                )
        assert resp.status_code == 200
        # Verify bucket_name was passed explicitly from settings
        put_call = mock_storage.put_object_bytes.call_args
        assert put_call.kwargs.get("bucket_name") == "my-staging-bucket"
        url_call = mock_storage.get_presigned_url.call_args
        assert url_call.kwargs.get("bucket_name") == "my-staging-bucket"
